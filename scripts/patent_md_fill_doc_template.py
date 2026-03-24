#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将「发明专利技术交底书」Markdown 按栏目填入 Word 模板（.doc 经 LibreOffice 转为 .docx 后处理）。

依赖: python-docx, LibreOffice (soffice)

一键模式（推荐，零参数）:
  python3 scripts/patent_md_fill_doc_template.py
  - 自动选择 patent_outputs 下最新的 .md
  - 自动使用仓库内模板 .cursor/skills/patent/软件实例-专利技术交底书模版.doc
  - 自动输出到 patent_outputs/filled-<md主名>-<时间>.doc

手动模式:
  python3 scripts/patent_md_fill_doc_template.py \\
    --md patent_outputs/xxx.md \\
    --template .cursor/skills/patent/软件实例-专利技术交底书模版.doc \\
    -o patent_outputs/xxx-filled.doc
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.text.paragraph import Paragraph
except ImportError:
    print("请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 模板中一级栏目标题使用全角点「1．」等，与填写说明「1.1」区分
SECTION_MARKERS = ("1．", "2．", "3．", "4．")
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
DEFAULT_TEMPLATE = (
    REPO_ROOT / ".cursor" / "skills" / "patent" / "软件实例-专利技术交底书模版.doc"
)
DEFAULT_OUTPUT_DIR = REPO_ROOT / "patent_outputs"


@dataclass
class MdBlock:
    kind: str  # "paragraph" | "bullet" | "image" | "subheading"
    text: str = ""
    alt: str = ""
    path: Optional[Path] = None


def _run_soffice_convert(src: Path, out_dir: Path, target_ext: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        "soffice",
        "--headless",
        "--convert-to",
        target_ext,
        "--outdir",
        str(out_dir),
        str(src),
    ]
    r = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if r.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换失败: {r.stderr or r.stdout}")
    base = src.stem + "." + target_ext
    out = out_dir / base
    if not out.is_file():
        raise FileNotFoundError(f"未找到转换输出: {out}")
    return out


def doc_to_docx(doc_path: Path, workdir: Path) -> Path:
    if doc_path.suffix.lower() == ".docx":
        return doc_path
    return _run_soffice_convert(doc_path, workdir, "docx")


def docx_to_doc(docx_path: Path, out_dir: Path) -> Path:
    return _run_soffice_convert(docx_path, out_dir, "doc")


def _para_full_text(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs if r.text)


def _is_template_section_title(text: str, section_num: int) -> bool:
    t = text.strip()
    prefix = f"{section_num}．"
    return t.startswith(prefix)


def _find_section_paragraphs(doc: Document) -> dict[int, Paragraph]:
    mapping: dict[int, Paragraph] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        for n in (1, 2, 3, 4):
            if _is_template_section_title(t, n):
                mapping[n] = p
                break
    for n in (1, 2, 3, 4):
        if n not in mapping:
            raise ValueError(f"模板中未找到第 {n} 节标题段落（应以「{n}．」开头）")
    return mapping


def _delete_following_until_next_section(
    header_p: Paragraph, next_section_num: Optional[int]
) -> None:
    el = header_p._element
    while True:
        nxt = el.getnext()
        if nxt is None:
            break
        tag = nxt.tag.split("}")[-1]
        if tag != "p":
            break
        texts = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
        if next_section_num is not None and _is_template_section_title(
            texts, next_section_num
        ):
            break
        parent = nxt.getparent()
        if parent is not None:
            parent.remove(nxt)


def insert_paragraph_after(
    paragraph: Paragraph, style: Optional[str] = None
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    return new_para


def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _parse_basic_info_table(md_chunk: str) -> dict[str, str]:
    rows: dict[str, str] = {}
    for line in md_chunk.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 2:
            continue
        key, val = cells[0], cells[1]
        if key in ("项目", "------", "---") or set(key) <= {"-", ":"}:
            continue
        rows[key] = val
    return rows


def _fill_info_table(doc: Document, info: dict[str, str]) -> None:
    if not doc.tables:
        return
    tbl = doc.tables[0]
    # 行: 发明设计名称 | 设计发明人/所属单位 | 电话号码/Email | 撰写日期
    key_to_cell = {
        "发明设计名称": (0, 1),
        "设计发明人": (1, 1),
        "所属单位": (1, 3),
        "电话号码": (2, 1),
        "Email": (2, 3),
        "撰写日期": (3, 1),
    }
    for key, (ri, ci) in key_to_cell.items():
        if key not in info:
            continue
        if ri < len(tbl.rows) and ci < len(tbl.rows[ri].cells):
            tbl.rows[ri].cells[ci].text = info[key]


def _resolve_image_path(md_path: Path, ref: str) -> Path:
    ref = ref.strip()
    p = (md_path.parent / ref).resolve()
    if p.is_file():
        return p
    alt = (md_path.parent / ref.lstrip("./")).resolve()
    if alt.is_file():
        return alt
    raise FileNotFoundError(f"图片不存在: {ref} (相对 {md_path.parent})")


_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _inline_images_to_blocks(text: str, md_path: Path) -> List[MdBlock]:
    """把纯文本段中的 ![alt](path) 拆成段落与图片块。"""
    blocks: List[MdBlock] = []
    rest = text.strip()
    while True:
        m = _IMAGE_RE.search(rest)
        if not m:
            if rest:
                blocks.append(MdBlock("paragraph", text=rest))
            break
        before = rest[: m.start()].strip()
        alt, path_s = m.group(1), m.group(2).strip()
        after = rest[m.end() :].strip()
        if before:
            blocks.append(MdBlock("paragraph", text=before))
        try:
            img_path = _resolve_image_path(md_path, path_s)
        except FileNotFoundError as e:
            blocks.append(
                MdBlock("paragraph", text=f"[图片缺失: {path_s}] ({e})")
            )
        else:
            blocks.append(MdBlock("image", alt=alt, path=img_path))
        rest = after
    return blocks


def _blocks_from_plain_text(text: str, md_path: Path) -> List[MdBlock]:
    """将一段不含列表/表格边界的正文转为 blocks（按 ### 拆小节，并处理内联图片）。"""
    blocks: List[MdBlock] = []
    text = text.strip()
    if not text or text == "---":
        return blocks
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("### "):
            blocks.append(MdBlock("subheading", text=lines[i].strip()[4:].strip()))
            i += 1
            seg_lines: List[str] = []
            while i < len(lines) and not lines[i].strip().startswith("### "):
                seg_lines.append(lines[i])
                i += 1
            seg = "\n".join(seg_lines).strip()
            if seg:
                blocks.extend(_inline_images_to_blocks(seg, md_path))
        else:
            seg_lines = []
            while i < len(lines) and not lines[i].strip().startswith("### "):
                seg_lines.append(lines[i])
                i += 1
            seg = "\n".join(seg_lines).strip()
            if seg:
                blocks.extend(_inline_images_to_blocks(seg, md_path))
    return blocks


def _md_body_to_blocks(body: str, md_path: Path) -> List[MdBlock]:
    blocks: List[MdBlock] = []
    lines = body.split("\n")
    i = 0
    buf: List[str] = []

    def flush_buf() -> None:
        nonlocal buf
        raw = "\n".join(buf).strip()
        buf = []
        if raw:
            blocks.extend(_blocks_from_plain_text(raw, md_path))

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_buf()
            i += 1
            chunk: List[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                chunk.append(lines[i])
                i += 1
            if chunk:
                blocks.append(MdBlock("paragraph", text="\n".join(chunk)))
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_buf()
            row_buf = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_buf.append(lines[i])
                i += 1
            table_md = "\n".join(row_buf)
            blocks.append(
                MdBlock(
                    "paragraph",
                    text="[表格已省略，请见 Markdown 源]\n" + table_md,
                )
            )
            continue
        if re.match(r"^[\s]*[-*]\s+", line):
            flush_buf()
            item = re.sub(r"^[\s]*[-*]\s+", "", line).strip()
            blocks.append(MdBlock("bullet", text=item))
            i += 1
            continue
        if stripped == "---":
            flush_buf()
            i += 1
            continue
        buf.append(line)
        i += 1
    flush_buf()
    return blocks


def _split_md_sections(md_text: str) -> List[Tuple[str, str]]:
    """返回 [(标题, 正文), ...]，标题不含 ## 前缀。"""
    sections: List[Tuple[str, str]] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []
    for line in md_text.splitlines():
        if line.startswith("## "):
            if current_title is not None:
                sections.append((current_title, "\n".join(current_lines).strip()))
            current_title = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, "\n".join(current_lines).strip()))
    return sections


def _map_md_to_section_num(title: str) -> Optional[int]:
    t = title.strip()
    if t.startswith("1．") or t.startswith("1."):
        if re.match(r"^1\.\d", t):
            return None
        return 1
    if t.startswith("2．") or t.startswith("2."):
        return 2
    if t.startswith("3．") or t.startswith("3."):
        if re.match(r"^3\.\d", t):
            return None
        return 3
    if t.startswith("4．") or t.startswith("4."):
        return 4
    return None


def _parse_md_file(md_path: Path) -> Tuple[dict[str, str], dict[int, str], str]:
    text = md_path.read_text(encoding="utf-8")
    sections = _split_md_sections(text)
    basic: dict[str, str] = {}
    bodies: dict[int, str] = {}
    appendix = ""
    for title, body in sections:
        if title == "基础信息":
            basic = _parse_basic_info_table(body)
            continue
        if title.startswith("待补"):
            appendix = body
            continue
        n = _map_md_to_section_num(title)
        if n is not None:
            bodies[n] = body
    return basic, bodies, appendix


def _append_blocks_after(
    anchor: Paragraph, blocks: List[MdBlock], picture_width_cm: float
) -> Paragraph:
    cur = anchor
    for b in blocks:
        if b.kind == "subheading":
            p = insert_paragraph_after(cur, style=None)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(b.text)
            r.bold = True
            r.font.size = Pt(12)
            cur = p
            continue
        if b.kind == "bullet":
            p = insert_paragraph_after(cur)
            try:
                p.style = "List"
            except KeyError:
                p.add_run("• ")
            _add_runs_with_bold(p, b.text)
            cur = p
            continue
        if b.kind == "paragraph":
            p = insert_paragraph_after(cur)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_runs_with_bold(p, b.text)
            cur = p
            continue
        if b.kind == "image":
            p = insert_paragraph_after(cur)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            try:
                run.add_picture(str(b.path), width=Cm(picture_width_cm))
            except Exception as e:
                p.add_run(f"[嵌入图片失败 {b.path}: {e}]")
            cur = p
            if b.alt:
                cap = insert_paragraph_after(cur)
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cap.add_run(b.alt).italic = True
                cur = cap
            continue
    return cur


def fill_template(
    template_docx: Path,
    md_path: Path,
    out_docx: Path,
    picture_width_cm: float = 14.0,
) -> None:
    basic, bodies, appendix = _parse_md_file(md_path)
    doc = Document(str(template_docx))
    _fill_info_table(doc, basic)

    sec_paras = _find_section_paragraphs(doc)
    # 从后往前删，避免 DOM 引用错位
    for n in (4, 3, 2, 1):
        header = sec_paras[n]
        next_n = n + 1 if n < 4 else None
        _delete_following_until_next_section(
            header, next_section_num=next_n if next_n else None
        )
    # 重新定位（删除后 Paragraph 包装可能失效）
    sec_paras = _find_section_paragraphs(doc)
    tail_after_sec4: Paragraph = sec_paras[4]
    for n in (1, 2, 3, 4):
        body = bodies.get(n, "")
        blocks = _md_body_to_blocks(body, md_path)
        end_p = _append_blocks_after(sec_paras[n], blocks, picture_width_cm)
        if n == 4:
            tail_after_sec4 = end_p
    if appendix.strip():
        extra = [MdBlock("subheading", text="待补字段清单")]
        extra.extend(_md_body_to_blocks(appendix, md_path))
        _append_blocks_after(tail_after_sec4, extra, picture_width_cm)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def _pick_latest_md(output_dir: Path) -> Path:
    cands = [p for p in output_dir.glob("*.md") if p.is_file()]
    if not cands:
        raise FileNotFoundError(f"未找到可用 Markdown 文件目录: {output_dir}")
    return max(cands, key=lambda p: p.stat().st_mtime)


def main() -> None:
    ap = argparse.ArgumentParser(description="Markdown 专利交底书填入 Word 模板")
    ap.add_argument(
        "--md",
        required=False,
        type=Path,
        help="交底书 .md 路径；不传则自动选择 ./patent_outputs 下最新 .md",
    )
    ap.add_argument(
        "--template",
        required=False,
        type=Path,
        help="模板 .doc 或 .docx；不传则使用仓库默认模板",
    )
    ap.add_argument(
        "-o",
        "--output",
        required=False,
        type=Path,
        help="输出 .docx 或 .doc；不传则自动写入 ./patent_outputs/filled-*.doc",
    )
    ap.add_argument(
        "--output-format",
        choices=("doc", "docx"),
        default="doc",
        help="当未指定 --output 时的默认输出格式，默认 doc",
    )
    ap.add_argument(
        "--picture-width-cm",
        type=float,
        default=14.0,
        help="正文插图宽度（厘米），默认 14",
    )
    ap.add_argument(
        "--keep-temp",
        action="store_true",
        help="保留临时目录（调试用）",
    )
    args = ap.parse_args()

    if args.md:
        md_path = args.md.expanduser().resolve()
    else:
        md_path = _pick_latest_md(DEFAULT_OUTPUT_DIR)

    if args.template:
        template_path = args.template.expanduser().resolve()
    else:
        template_path = DEFAULT_TEMPLATE.resolve()

    if args.output:
        out_path = args.output.expanduser().resolve()
    else:
        stamp = datetime.now().strftime("%Y%m%d-%H%M")
        stem = md_path.stem[:80]
        suffix = ".doc" if args.output_format == "doc" else ".docx"
        out_path = (DEFAULT_OUTPUT_DIR / f"filled-{stem}-{stamp}{suffix}").resolve()

    if not md_path.is_file():
        sys.exit(f"找不到 Markdown: {md_path}")
    if not template_path.is_file():
        sys.exit(f"找不到模板: {template_path}")

    tmp = Path(tempfile.mkdtemp(prefix="patent_md_fill_"))
    try:
        tpl_docx = doc_to_docx(template_path, tmp)
        work_docx = tmp / "work.docx"
        shutil.copy2(tpl_docx, work_docx)

        if out_path.suffix.lower() == ".docx":
            fill_template(work_docx, md_path, out_path, args.picture_width_cm)
        else:
            filled = tmp / "filled.docx"
            fill_template(work_docx, md_path, filled, args.picture_width_cm)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            converted = docx_to_doc(filled, out_path.parent)
            if converted.resolve() != out_path.resolve():
                shutil.move(str(converted), str(out_path))
    finally:
        if not args.keep_temp:
            shutil.rmtree(tmp, ignore_errors=True)

    print(f"已写入: {out_path}")


if __name__ == "__main__":
    main()
